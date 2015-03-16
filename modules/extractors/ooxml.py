#coding=utf-8
"""This module includes tools for handling and extracting text from
    Office Open XML and modern Mocrosoft Word (.docx) files.
"""

from modules.extractors.documentBase import ExtractorBase, Page

from docx import Document
import openxmllib  # good for metadata, bad for text
from modules.metadata import Metadata


class DocxPage(Page):

    def __init__(self):
        pass

    def get_text(self):
        """Returns all the text in one single page.
           We might be able to use e.g. Abiword to calculate
           approximate page breaks.
        """
        return self._text


class DocxExtractor(ExtractorBase):
    """Class for getting plain text from a Office Open XML file.
    """

    def get_metadata(self):
        """Returns a .modules.metadata.Metadata object
        """
        self.metadata = Metadata()
        document = openxmllib.openXmlDocument(path=self.path)
        self.metadata.add(document.allProperties, "ooxml")
        return self.metadata

    def get_header(self):
        """Our docx library has no support for headers yet.
           For now carve out the header ourselves by parsing xml files.
        """
        headers = []
        namespaces = dict(
            w="http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            v="urn:schemas-microsoft-com:vml",
            r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        )

        import xml.etree.ElementTree as ET

        """ Docx is basically a zip file with embedded xml and media files
        """
        import zipfile
        z = zipfile.ZipFile(self.path)

        """ Get all header files and parse them for text
        """
        docx_files = z.namelist()
        header_xml_files = filter(lambda x: x.startswith('word/header'),
                                  docx_files)

        for header_xml in header_xml_files:
            # header_name = header_xml.replace('word/', '')
            xml_file = z.open(header_xml, 'r')
            root = ET.fromstring(xml_file.read())

            text_elements = root.findall(".//w:t", namespaces)
            for text_element in text_elements:
                if text_element.text is not None:
                    headers.append(text_element.text)

            """ Headers might also include images with text
                To get those we find all images in header
                    => get id's of those images
                    => find corresponding image files in the /media directory
                    => OCR
            """
            """
            # Disiable OCR, as PIL on *nix can't handle WMF images
            images = root.findall(".//v:imagedata", namespaces)
            for image in images:
                id_key = '{%s}id' % namespaces['r']
                image_id = image.attrib[id_key]
                relation_file_path = 'word/_rels/%s.rels' % header_name
                relation_file = z.open(relation_file_path, 'r')
                relation_file_root = ET.fromstring(relation_file.read())
                xpath = ".//*[@Id='%s']" % image_id
                image_path = 'word/%s' % relation_file_root.find(xpath, namespaces).attrib['Target']

                # OCR
                import pytesseract
                from PIL import Image
                import cStringIO
                img = Image.open(cStringIO.StringIO(z.open(image_path).read()))
                string_from_image = pytesseract.image_to_string(img, lang='swe')
                headers.append(string_from_image.decode('utf-8'))
            """
        return "\n".join(headers)


    def get_next_page(self):
        """Returns all the text in one single page.
           We might be able to use e.g. Abiword to calculate
           approximate page breaks.
        """
        page = DocxPage()
        page._text = self.get_text()
        yield page

    def get_text(self):
        """Returns all text content from the document as plain text.
        """
        try:
            return self._text_cache
        except AttributeError:  # not cached
            pass

        paratextlist = []
        document = Document(self.path)
        for paragraph in document.paragraphs:
            paratextlist.append(paragraph.text)
        self._text_cache = "\n".join(paratextlist)
        return self._text_cache

if __name__ == "__main__":
    print "This module is only intended to be called from other scripts."
    import sys
    sys.exit()
